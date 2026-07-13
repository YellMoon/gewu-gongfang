from __future__ import annotations

import json
import os
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path


PARSER_DIR = Path(__file__).resolve().parents[1]
if str(PARSER_DIR) not in sys.path:
    sys.path.insert(0, str(PARSER_DIR))

from docx_fixture import DocxFixture  # noqa: E402
from parse_word import _image_tag, build_question_rich_content, quality_report  # noqa: E402


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
O = "urn:schemas-microsoft-com:office:office"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


DOCUMENT = f'''<w:document xmlns:w="{W}" xmlns:m="{M}"><w:body><w:p>
  <w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr>
  <w:r><w:t>Find the value </w:t></w:r>
  <m:oMath><m:f><m:num><m:r><m:t>a</m:t></m:r></m:num><m:den><m:r><m:t>b</m:t></m:r></m:den></m:f></m:oMath>
  <w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText> EQ \\r(x) </w:instrText></w:r>
  <w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>sqrt x</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r>
</w:p></w:body></w:document>'''

NUMBERING = f'''<w:numbering xmlns:w="{W}">
  <w:abstractNum w:abstractNumId="0"><w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/></w:lvl></w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
</w:numbering>'''


def _formula_node_ids(document):
    return [node["attrs"]["id"] for paragraph in document.get("content", []) for node in paragraph.get("content", []) if node.get("type") == "formula"]


class ParseWordFormulaIntegrationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.fixture = DocxFixture().add("word/document.xml", DOCUMENT).add("word/numbering.xml", NUMBERING)
        self.path = self.fixture.write()

    def tearDown(self) -> None:
        self.fixture.cleanup()

    def test_forced_xml_fallback_returns_canonical_formulas_and_quality_counts(self):
        env = dict(os.environ)
        env["GEWU_FORCE_DOCX_XML_FALLBACK"] = "1"
        completed = subprocess.run(
            [sys.executable, str(PARSER_DIR / "parse_word.py"), str(self.path), "lecture"],
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            env=env,
            timeout=30,
        )
        result = json.loads(completed.stdout)
        self.assertEqual(result["count"], 1)
        formulas = result["questions"][0]["formulas"]
        self.assertEqual([item["canonical_latex"] for item in formulas], [r"\frac{a}{b}", r"\sqrt{x}"])
        self.assertEqual([item["source"]["source_format"] for item in formulas], ["omml", "eq_field"])
        self.assertEqual(result["quality_report"]["formula_import"]["total"], 2)
        self.assertEqual(result["quality_report"]["formula_import"]["needs_review"], 0)
        rich = result["questions"][0]["rich_content"]
        self.assertEqual(rich["type"], "question-document")
        nodes = rich["sections"]["stem"]["content"][0]["content"]
        self.assertEqual([node["attrs"]["canonicalLatex"] for node in nodes if node["type"] == "formula"], [r"\frac{a}{b}", r"\sqrt{x}"])

    def test_image_tag_escapes_malicious_file_name_attribute(self):
        tag = _image_tag({"content_hash": "safe", "file_name": 'x" onerror="alert(1).png'})
        self.assertIn('alt="x&amp;quot; onerror=&amp;quot;alert(1).png"'.replace("&amp;quot;", "&quot;"), tag)
        self.assertNotIn(' onerror="', tag)

    def test_python_docx_path_returns_token_derived_rich_formula(self):
        try:
            from docx import Document
            from docx.oxml import OxmlElement, parse_xml
            from docx.oxml.ns import qn
        except ImportError:
            self.skipTest("python-docx is not installed")
        descriptor, raw_path = tempfile.mkstemp(suffix=".docx")
        os.close(descriptor)
        path = Path(raw_path)
        try:
            document = Document()
            paragraph = document.add_paragraph(style="List Number")
            num_pr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "0")
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), "5")
            num_pr.extend([ilvl, num_id])
            paragraph._p.get_or_add_pPr().append(num_pr)
            paragraph.add_run("Find ")
            paragraph._p.append(parse_xml(f'<m:oMath xmlns:m="{M}"><m:sSup><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup></m:oMath>'))
            document.save(path)
            completed = subprocess.run(
                [sys.executable, str(PARSER_DIR / "parse_word.py"), str(path), "lecture"],
                check=True,
                capture_output=True,
                text=True,
                encoding="utf-8",
                timeout=30,
            )
            result = json.loads(completed.stdout)
            self.assertEqual(result["count"], 1)
            self.assertEqual(result["questions"][0]["formulas"][0]["canonical_latex"], "x^{2}")
            nodes = result["questions"][0]["rich_content"]["sections"]["stem"]["content"][0]["content"]
            self.assertEqual([node["attrs"]["canonicalLatex"] for node in nodes if node["type"] == "formula"], ["x^{2}"])
        finally:
            path.unlink(missing_ok=True)

    def test_comment_formula_is_canonical_answer_node_not_body_source(self):
        question = {"stem": "body", "answer": "comment", "options": [], "sub_questions": [], "formulas": [
            {"canonical_latex": r"x^2", "source": {"part_name": "word/document.xml"}},
            {"canonical_latex": r"\sqrt{y}", "source": {"part_name": "word/comments.xml"}},
        ]}
        rich = build_question_rich_content(question)
        stem_nodes = rich["sections"]["stem"]["content"][0]["content"]
        answer_nodes = rich["sections"]["answer"]["content"][0]["content"]
        self.assertEqual([node["attrs"]["canonicalLatex"] for node in stem_nodes if node["type"] == "formula"], [r"x^2"])
        self.assertEqual([node["attrs"]["canonicalLatex"] for node in answer_nodes if node["type"] == "formula"], [r"\sqrt{y}"])

    def test_parser_optional_null_attrs_match_rich_content_normalizer_contract(self):
        rich = build_question_rich_content({
            "stem": '<span data-formula-id="f-null" data-latex="x"></span><img src="question-asset://asset-null" alt="diagram">',
            "answer": "", "analysis": "", "options": [], "sub_questions": [],
        })
        nodes = rich["sections"]["stem"]["content"][0]["content"]
        formula = next(node for node in nodes if node["type"] == "formula")
        image = next(node for node in nodes if node["type"] == "image")
        self.assertIsNone(formula["attrs"]["sourceRef"])
        self.assertIsNone(image["attrs"]["width"])

    def test_rich_content_preserves_formula_field_ownership_and_source_order(self):
        formula = lambda latex, paragraph, content: {
            "id": f"formula-{paragraph}-{content}",
            "canonical_latex": latex,
            "conversion_status": "complete",
            "source": {"source_format": "omml", "part_name": "word/document.xml", "paragraph_index": paragraph, "content_index": content},
        }
        question = {
            "stem": 'stem <span data-formula-id="formula-0-1" data-latex="s"></span>',
            "options": [{"label": "A", "content": 'choice <span data-formula-id="formula-1-1" data-latex="o"></span>'}],
            "sub_questions": [{"title": "(1)", "content": 'part <span data-formula-id="formula-2-1" data-latex="q"></span>', "answer": 'subanswer <span data-formula-id="formula-3-1" data-latex="u"></span>'}],
            "answer": 'answer <span data-formula-id="formula-4-1" data-latex="a"></span>',
            "analysis": 'analysis <span data-formula-id="formula-5-1" data-latex="e"></span>',
            "formulas": [formula("s", 0, 1), formula("o", 1, 1), formula("q", 2, 1), formula("u", 3, 1), formula("a", 4, 1), formula("e", 5, 1)],
        }

        rich = build_question_rich_content(question)

        def latex(document):
            return [node["attrs"]["canonicalLatex"] for paragraph in document["content"] for node in paragraph.get("content", []) if node["type"] == "formula"]

        self.assertEqual(latex(rich["sections"]["stem"]), ["s"])
        self.assertEqual(latex(rich["sections"]["options"][0]["content"]), ["o"])
        self.assertEqual(latex(rich["sections"]["subQuestions"][0]["content"]), ["q"])
        self.assertEqual(latex(rich["sections"]["subQuestions"][0]["answer"]), ["u"])
        self.assertEqual(latex(rich["sections"]["answer"]), ["a"])
        self.assertEqual(latex(rich["sections"]["analysis"]), ["e"])

    def test_quality_report_lists_actionable_formula_locations_by_source_and_status(self):
        question = {"stem": "Q", "answer": "A", "formulas": [{
            "id": "formula-review",
            "canonical_latex": "x",
            "conversion_status": "approximate",
            "warnings": ["unsupported construct"],
            "source": {"source_format": "eq_field", "part_name": "word/document.xml", "paragraph_index": 7, "content_index": 3},
        }]}

        report = quality_report([question])["formula_import"]

        self.assertEqual(report["by_source"], {"eq_field": 1})
        self.assertEqual(report["by_status"], {"approximate": 1})
        self.assertEqual(report["issues"], [{
            "formula_id": "formula-review",
            "status": "approximate",
            "source_format": "eq_field",
            "location": "word/document.xml:paragraph[7]:content[3]:question[index=0]:field[unknown]",
            "question_index": 0,
            "question_number": None,
            "field": "unknown",
            "warnings": ["unsupported construct"],
        }])

    def test_formula_markup_preserves_duplicate_latex_occurrences_by_formula_id(self):
        question = {
            "stem": 'left <span data-formula-id="f1" data-latex="x"></span> middle <span data-formula-id="f2" data-latex="x"></span> right',
            "options": [{"label": "A", "content": '<span data-formula-id="f3" data-latex="x"></span>'}],
            "sub_questions": [], "answer": "", "analysis": "",
            "formulas": [
                {"id": "f1", "canonical_latex": "x", "source": {"part_name": "word/document.xml"}},
                {"id": "f2", "canonical_latex": "x", "source": {"part_name": "word/document.xml"}},
                {"id": "f3", "canonical_latex": "x", "source": {"part_name": "word/document.xml"}},
            ],
        }
        rich = build_question_rich_content(question)
        stem = rich["sections"]["stem"]["content"][0]["content"]
        option = rich["sections"]["options"][0]["content"]["content"][0]["content"]
        self.assertEqual([node["attrs"]["id"] for node in stem if node["type"] == "formula"], ["f1", "f2"])
        self.assertEqual([node["attrs"]["id"] for node in option if node["type"] == "formula"], ["f3"])
        review_question = {**question, "formulas": [
            {**question["formulas"][2], "conversion_status": "approximate", "warnings": []},
            {"id": "fc", "canonical_latex": "c", "conversion_status": "approximate", "warnings": [], "source": {"source_format": "omml", "part_name": "word/comments.xml", "comment_id": "9", "paragraph_index": 0, "content_index": 1}},
        ], "answer": '<span data-formula-id="fc" data-latex="c"></span>'}
        issues = quality_report([review_question])["formula_import"]["issues"]
        self.assertEqual([item["field"] for item in issues], ["option[0]", "answer"])
        self.assertIn("comment[9]", issues[1]["location"])

    def test_exam_fixture_attaches_answer_and_analysis_formula_metadata(self):
        frac = '<m:oMath><m:f><m:num><m:r><m:t>x</m:t></m:r></m:num><m:den><m:r><m:t>y</m:t></m:r></m:den></m:f></m:oMath>'
        eq = '<w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText> EQ \\r(z) </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>sqrt z</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r>'
        document = f'''<w:document xmlns:w="{W}" xmlns:m="{M}"><w:body>
          <w:p><w:r><w:t>1. Stem </w:t></w:r>{frac}</w:p>
          <w:p><w:r><w:t>A. Option </w:t></w:r>{eq}</w:p>
          <w:p><w:r><w:t>(1) Part </w:t></w:r>{frac}</w:p>
          <w:p><w:r><w:t>参考答案</w:t></w:r></w:p>
          <w:p><w:r><w:t>1. Answer </w:t></w:r>{frac}</w:p>
          <w:p><w:r><w:t>(1) Subanswer </w:t></w:r>{frac}</w:p>
          <w:p><w:r><w:t>&#12304;&#35299;&#26512;&#12305;Reason </w:t></w:r>{eq}</w:p>
        </w:body></w:document>'''
        fixture = DocxFixture().add("word/document.xml", document)
        path = fixture.write()
        try:
            env = dict(os.environ, GEWU_FORCE_DOCX_XML_FALLBACK="1")
            completed = subprocess.run([sys.executable, str(PARSER_DIR / "parse_word.py"), str(path), "exam"], check=True, capture_output=True, text=True, encoding="utf-8", env=env, timeout=30)
            result = json.loads(completed.stdout)
            question = result["questions"][0]
            self.assertEqual(len(question["formulas"]), 6)
            self.assertEqual(result["quality_report"]["formula_import"]["total"], 6)
            answer_ids = _formula_node_ids(question["rich_content"]["sections"]["answer"])
            analysis_ids = _formula_node_ids(question["rich_content"]["sections"]["analysis"])
            self.assertEqual(answer_ids, [question["formulas"][3]["id"]])
            subanswer_ids = _formula_node_ids(question["rich_content"]["sections"]["subQuestions"][0]["answer"])
            self.assertEqual(subanswer_ids, [question["formulas"][4]["id"]])
            self.assertEqual(analysis_ids, [question["formulas"][5]["id"]])
            review = quality_report([{**question, "formulas": [{**question["formulas"][4], "conversion_status": "approximate"}]}])["formula_import"]["issues"][0]
            self.assertEqual(review["field"], "subanswer[0]")
        finally:
            fixture.cleanup()

    def test_lecture_fixture_preserves_comment_and_table_formula_coordinates(self):
        document = f'''<w:document xmlns:w="{W}" xmlns:m="{M}" xmlns:o="{O}" xmlns:r="{R}" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><w:body>
          <w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr><w:commentRangeStart w:id="7"/><w:r><w:t>Stem </w:t></w:r><m:oMath><m:r><m:t>x</m:t></m:r></m:oMath><w:commentRangeEnd w:id="7"/><w:r><w:commentReference w:id="7"/></w:r></w:p>
          <w:p><w:r><w:t>A. Option </w:t></w:r><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText> EQ \\r(y) </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>sqrt y</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>
          <w:tbl><w:tr><w:tc><w:p><w:r><w:t>Cell </w:t></w:r><w:r><w:object><o:OLEObject ProgID="Equation.DSMT4" r:id="rOle"/></w:object><w:drawing><a:blip r:embed="rImage"/></w:drawing></w:r></w:p></w:tc></w:tr></w:tbl>
        </w:body></w:document>'''
        comments = f'''<w:comments xmlns:w="{W}" xmlns:m="{M}"><w:comment w:id="7"><w:p><w:r><w:t>Answer </w:t></w:r><m:oMath><m:r><m:t>c</m:t></m:r></m:oMath></w:p></w:comment></w:comments>'''
        rels = '''<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rOle" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" Target="embeddings/equation.bin"/><Relationship Id="rImage" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/equation.png"/></Relationships>'''
        fixture = (DocxFixture().add("word/document.xml", document).add("word/numbering.xml", NUMBERING).add("word/comments.xml", comments).add("word/_rels/document.xml.rels", rels).add("word/embeddings/equation.bin", b"OLE-STUB").add("word/media/equation.png", b"PNG-STUB"))
        path = fixture.write()
        try:
            env = dict(os.environ, GEWU_FORCE_DOCX_XML_FALLBACK="1")
            completed = subprocess.run([sys.executable, str(PARSER_DIR / "parse_word.py"), str(path), "lecture"], check=True, capture_output=True, text=True, encoding="utf-8", env=env, timeout=30)
            result = json.loads(completed.stdout)
            question = result["questions"][0]
            self.assertEqual(result["quality_report"]["formula_import"]["by_source"], {"omml": 2, "eq_field": 1, "mathtype": 1})
            table_formula = next(item for item in question["formulas"] if item["source"]["source_format"] == "mathtype")
            self.assertEqual((table_formula["source"]["table_row"], table_formula["source"]["table_cell"], table_formula["source"]["cell_paragraph"]), (0, 0, 0))
            stem_nodes = [node for paragraph in question["rich_content"]["sections"]["stem"]["content"] for node in paragraph.get("content", [])]
            ole_node_index = next(index for index, node in enumerate(stem_nodes) if node.get("type") == "formula" and node.get("attrs", {}).get("id") == table_formula["id"])
            ole_node = stem_nodes[ole_node_index]
            self.assertEqual(ole_node["attrs"]["conversionStatus"], "preview_only")
            self.assertEqual(ole_node["attrs"]["sourceFormat"], "mathtype")
            self.assertEqual(ole_node["attrs"]["previewRef"], "word/media/equation.png")
            self.assertTrue(question["assets"])
            self.assertTrue(any(node.get("type") == "image" for node in stem_nodes[ole_node_index + 1:]))
            comment_formula = next(item for item in question["formulas"] if item["source"].get("comment_id") == "7")
            self.assertIn(comment_formula["id"], _formula_node_ids(question["rich_content"]["sections"]["answer"]))
            issue = next(item for item in result["quality_report"]["formula_import"]["issues"] if item["formula_id"] == table_formula["id"])
            self.assertIn("table[row=0,cell=0,paragraph=0]", issue["location"])
            self.assertEqual(issue["question_index"], 0)
        finally:
            fixture.cleanup()


if __name__ == "__main__":
    unittest.main()
