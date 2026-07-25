from __future__ import annotations
import hashlib
import json
import sys
import zipfile
from pathlib import Path
from docx import Document

def paragraph_data(paragraph, index):
    return {
        "index": index,
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": int(paragraph.alignment) if paragraph.alignment is not None else None,
        "runs": [{"text": run.text, "font": run.font.name, "size_pt": run.font.size.pt if run.font.size else None, "bold": run.bold, "italic": run.italic} for run in paragraph.runs],
    }

def main():
    source = Path(sys.argv[1]).resolve()
    doc = Document(str(source))
    with zipfile.ZipFile(source) as archive:
        parts = [{"path": item.filename, "size": item.file_size, "sha256": hashlib.sha256(archive.read(item.filename)).hexdigest()} for item in archive.infolist() if not item.is_dir()]
    result = {
        "source": str(source), "size": source.stat().st_size, "sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
        "paragraphs": [paragraph_data(paragraph, index) for index, paragraph in enumerate(doc.paragraphs)],
        "tables": [{"index": index, "rows": len(table.rows), "columns": len(table.columns), "cells": [[cell.text for cell in row.cells] for row in table.rows]} for index, table in enumerate(doc.tables)],
        "sections": [{"index": index, "header": [paragraph.text for paragraph in section.header.paragraphs], "footer": [paragraph.text for paragraph in section.footer.paragraphs], "first_header": [paragraph.text for paragraph in section.first_page_header.paragraphs], "first_footer": [paragraph.text for paragraph in section.first_page_footer.paragraphs]} for index, section in enumerate(doc.sections)],
        "parts": parts,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    main()
